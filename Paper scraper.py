import os
import random
import re
import time
from flask import Flask, render_template, request, send_file
from bs4 import BeautifulSoup
import requests
from docx import Document

app = Flask(__name__)

# User agents list for random selection
USER_AGENTS = [
    'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36',
    'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36',
    # Add more user agents as needed
]

# Function to fetch URL content
def fetch_url(url, session):
    headers = {
        'User-Agent': random.choice(USER_AGENTS),
        'Referer': 'https://www.google.com',
    }
    try:
        response = session.get(url, headers=headers)
        response.raise_for_status()  # Raise an HTTPError for bad responses
        return response.content, None
    except requests.HTTPError as e:
        if e.response.status_code == 403:
            return None, f"HTTP error 403: Forbidden when accessing {url}"
        else:
            return None, f"HTTP error {e.response.status_code} when accessing {url}"
    except requests.RequestException as e:
        return None, f"Request error: {e} when accessing {url}"
    except Exception as e:
        return None, f"An error occurred: {e}"

# Function to find DOI links in HTML content
def find_doi_links(html_content):
    if not html_content:
        return []

    soup = BeautifulSoup(html_content, 'html.parser')
    doi_pattern = re.compile(r'https://doi\.org/([^"\s]+)')
    doi_links = []

    tags = soup.find_all('a', href=True)
    for tag in tags:
        href = tag.get('href')
        if doi_pattern.match(href):
            doi_links.append(href)

    return doi_links

# Function to create a Word document with the results
def create_word_document(main_doi_links, subpage_dois, error_urls):
    doc = Document()
    doc.add_heading('DOI Extraction Results', level=1)

    if main_doi_links:
        doc.add_heading('DOI Links Found on Main Page:', level=2)
        for doi in main_doi_links:
            doc.add_paragraph(doi, style='List Bullet')

    if subpage_dois:
        doc.add_heading('DOI Links Found on Subpages:', level=2)
        for link, dois in subpage_dois.items():
            doc.add_heading(link, level=3)
            for doi in dois:
                doc.add_paragraph(doi, style='List Bullet')

    if error_urls:
        doc.add_heading('Errors Encountered:', level=2)
        for url, error in error_urls:
            doc.add_paragraph(f'{url}: {error}', style='List Bullet')

    file_path = os.path.join('results', 'doi_results.docx')
    doc.save(file_path)
    return file_path

@app.route('/', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':
        url = request.form['url']
        session = requests.Session()
        error_urls = []
        main_doi_links = []
        subpage_dois = {}

        try:
            html, error_message = fetch_url(url, session)

            if error_message:
                error_urls.append((url, error_message))
            else:
                main_doi_links.extend(find_doi_links(html))
                soup = BeautifulSoup(html, 'html.parser')

                if not main_doi_links:
                    tags = soup.find_all('a', href=True)
                    for tag in tags:
                        link = tag.get('href')
                        if not link.startswith(('http://', 'https://')):
                            continue
                        try:
                            time.sleep(random.uniform(1, 3))
                            sub_html, sub_error_message = fetch_url(link, session)

                            if sub_error_message:
                                error_urls.append((link, sub_error_message))
                            else:
                                sub_doi_links = find_doi_links(sub_html)
                                if sub_doi_links:
                                    subpage_dois[link] = sub_doi_links

                        except requests.RequestException as e:
                            error_urls.append((link, str(e)))

        except requests.RequestException as e:
            error_urls.append((url, str(e)))

        # Create a Word document with the results
        file_path = create_word_document(main_doi_links, subpage_dois, error_urls)
        return render_template('result.html', main_doi_links=main_doi_links, subpage_dois=subpage_dois, error_urls=error_urls, file_path=file_path)

    return render_template('index.html')

@app.route('/download/<filename>')
def download_file(filename):
    return send_file(os.path.join('results', filename), as_attachment=True)

if __name__ == '__main__':
    if not os.path.exists('results'):
        os.makedirs('results')
    app.run(debug=True)
